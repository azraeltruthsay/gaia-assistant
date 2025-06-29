import os
import logging
from typing import Optional, List
from datetime import datetime

from striprtf.striprtf import rtf_to_text
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_community.document_loaders import TextLoader

logger = logging.getLogger("GAIA")

class DocumentProcessor:
    """
    Handles loading, preprocessing, and converting documents into structured markdown or LangChain document objects.
    Used for both initial data ingestion and markdown embedding during boot.
    """

    def __init__(self, config, llm=None):
        """
        Initialize with configuration and optional LLM for markdown conversion.

        Args:
            config: Configuration object
            llm: Optional language model for structured conversion
        """
        self.config = config
        self.llm = llm

    def extract_text_from_file(self, filepath: str) -> Optional[str]:
        """Extract text content from .txt, .md, .docx, or .rtf files."""
        _, ext = os.path.splitext(filepath)
        try:
            if ext.lower() in [".txt", ".md"]:
                with open(filepath, 'r', encoding='utf-8') as f:
                    return f.read()
            elif ext.lower() == ".rtf":
                return self._extract_rtf(filepath)
            elif ext.lower() == ".docx":
                return self._extract_docx(filepath)
            else:
                logger.warning(f"Unsupported file format: {ext}")
                return None
        except Exception as e:
            logger.error(f"Error extracting text from {filepath}: {e}")
            return None

    def _extract_rtf(self, filepath: str) -> Optional[str]:
        """Extract text from RTF files with utf-8 and latin-1 fallback."""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return rtf_to_text(f.read())
        except UnicodeDecodeError:
            try:
                with open(filepath, 'r', encoding='latin-1') as f:
                    return rtf_to_text(f.read())
            except Exception as e:
                logger.error(f"Error extracting RTF with latin-1 encoding: {e}")
                return None
        except Exception as e:
            logger.error(f"Error extracting RTF with utf-8 encoding: {e}")
            return None

    def _extract_docx(self, filepath: str) -> Optional[str]:
        """Extract text from Microsoft Word .docx files."""
        try:
            doc = DocxDocument(filepath)
            return "\n".join(para.text for para in doc.paragraphs)
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            return None

    def convert_to_markdown(self, text: str) -> Optional[str]:
        """
        Convert plain text to structured markdown using the LLM.

        Args:
            text: Raw extracted text

        Returns:
            str or None: Markdown string, or None on failure
        """
        if not self.llm:
            logger.error("LLM not available for markdown conversion")
            return None

        prompt = f"""Convert the following text into a structured markdown document:

{text}

Create a logical organization with appropriate headers and formatting. 
Use # for primary headers, ## for secondary headers, etc.
Avoid including embedded HTML, links, or code blocks unless the text requires it.

Markdown Output:
"""
        try:
            result = self.llm(prompt)
            if not result or len(result.strip()) < 10:
                logger.warning("⚠️ LLM returned empty or too-short markdown")
                return None

            lines = result.strip().splitlines()
            cleaned = []
            seen_headers = set()
            for line in lines:
                if line.strip().lower() in seen_headers:
                    continue
                if line.startswith("<") and line.endswith(">"):
                    continue  # strip embedded HTML
                cleaned.append(line.strip())
                seen_headers.add(line.strip().lower())

            markdown_output = "\n".join(cleaned)
            return markdown_output
        except Exception as e:
            logger.error(f"Error during markdown conversion: {e}")
            return None

    def save_markdown(self, filepath: str, content: str) -> bool:
        """Write markdown content to file."""
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            logger.info(f"📄 Saved markdown to {filepath}")
            return True
        except Exception as e:
            logger.error(f"Error saving markdown to {filepath}: {e}")
            return False

    def load_and_preprocess_data(self, data_path: str) -> List[Document]:
        """
        Load and preprocess markdown documents from a directory for embedding.

        Args:
            data_path: Directory containing markdown files

        Returns:
            List[Document]: LangChain documents
        """
        documents = []
        try:
            for filename in os.listdir(data_path):
                if filename.endswith(".md"):
                    filepath = os.path.join(data_path, filename)
                    try:
                        loader = TextLoader(filepath, encoding='utf-8')
                        documents.extend(loader.load())
                        logger.info(f"📑 Loaded document: {filepath}")
                    except Exception as e:
                        logger.error(f"Error loading document {filepath}: {e}")
        except Exception as e:
            logger.error(f"Error accessing directory {data_path}: {e}")

        return documents

    def process_raw_data(self) -> None:
        """Process raw input files in `raw_data_path` into markdown documents."""
        if not self.llm:
            logger.error("LLM not available for raw data processing")
            return

        try:
            for filename in os.listdir(self.config.raw_data_path):
                filepath = os.path.join(self.config.raw_data_path, filename)
                logger.info(f"📥 Processing raw file: {filepath}")

                raw_text = self.extract_text_from_file(filepath)
                if not raw_text:
                    logger.warning(f"⚠️ Could not extract text from {filepath}")
                    continue

                markdown_content = self.convert_to_markdown(raw_text)
                if not markdown_content:
                    logger.warning(f"⚠️ Could not convert {filename} to markdown")
                    continue

                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_filename = f"converted_{os.path.splitext(filename)[0]}_{timestamp}.md"
                output_filepath = os.path.join(self.config.output_path, output_filename)

                self.save_markdown(output_filepath, markdown_content)
        except Exception as e:
            logger.error(f"Error processing raw data: {e}")

    def get_document_info(self, filepath: str):
        """Return basic file metadata for a given document."""
        try:
            filename = os.path.basename(filepath)
            mtime = os.path.getmtime(filepath)
            size = os.path.getsize(filepath)
            return {
                'name': filename,
                'path': filepath,
                'modified': datetime.fromtimestamp(mtime).strftime('%Y-%m-%d %H:%M:%S'),
                'size': size
            }
        except Exception as e:
            logger.error(f"Error getting document info for {filepath}: {e}")
            return None

    def process_documents(self, directory: str, tier: Optional[str] = None, project: Optional[str] = None) -> List[Document]:
        """Load and wrap markdown documents from a directory with metadata."""
        documents = []
        if not os.path.isdir(directory):
            logger.warning(f"Directory not found or invalid: {directory}")
            return documents

        try:
            for filename in os.listdir(directory):
                if filename.endswith(".md"):
                    filepath = os.path.join(directory, filename)
                    try:
                        with open(filepath, 'r', encoding='utf-8') as f:
                            content = f.read()
                            if not content.strip():
                                logger.warning(f"Empty file skipped: {filepath}")
                                continue

                            metadata = {
                                "filename": filename,
                                "source_path": filepath,
                                "tier": tier or "unspecified",
                                "project": project or "global"
                            }
                            documents.append(Document(page_content=content, metadata=metadata))
                            logger.info(f"📘 Loaded document with metadata: {metadata}")
                    except Exception as e:
                        logger.error(f"Failed to load file {filepath}: {e}")
        except Exception as e:
            logger.error(f"Error accessing directory {directory}: {e}")

        return documents

    def embed_documents(self) -> int:
        """Reads all markdown files from `structured/`, embeds them via the vector store."""
        path = self.config.structured_path
        documents = self.process_documents(path, tier="2_semantic")
        if not documents:
            return 0
        self.vector_store_manager.add_documents(documents)
        return len(documents)

    def generate_artifacts(self) -> int:
        """Processes raw documents and converts them into markdown in the output directory."""
        self.process_raw_data()
        files = os.listdir(self.config.output_path)
        return len([f for f in files if f.endswith(".md")])
        """
        Load and wrap markdown documents from a directory with metadata.

        Args:
            directory: Path to files
            tier: Knowledge tier identifier (e.g., '0_system_reference')
            project: Optional project name to associate

        Returns:
            List[Document]: LangChain documents with metadata
        """
        documents = []
        if not os.path.isdir(directory):
            logger.warning(f"Directory not found or invalid: {directory}")
            return documents

        try:
            for filename in os.listdir(directory):
                if filename.endswith(".md"):
                    filepath = os.path.join(directory, filename)
                    try:
                        with open(filepath, 'r', encoding='utf-8') as f:
                            content = f.read()
                            if not content.strip():
                                logger.warning(f"Empty file skipped: {filepath}")
                                continue

                            metadata = {
                                "filename": filename,
                                "source_path": filepath,
                                "tier": tier or "unspecified",
                                "project": project or "global"
                            }
                            documents.append(Document(page_content=content, metadata=metadata))
                            logger.info(f"📘 Loaded document with metadata: {metadata}")
                    except Exception as e:
                        logger.error(f"Failed to load file {filepath}: {e}")
        except Exception as e:
            logger.error(f"Error accessing directory {directory}: {e}")

        return documents
